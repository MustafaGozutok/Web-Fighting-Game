#!/usr/bin/env python3
"""Generate Software Requirements Document for Color Clash in DOCX format."""

import subprocess
import sys

# Install python-docx if needed
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

# ══════════════════════ TITLE PAGE ══════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Software Requirements Document')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Color Clash – A Two-Player Fighting Game')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('TOBB ETU – BIL 482\nFebruary 2026')
run.font.size = Pt(12)

doc.add_paragraph()
team = doc.add_paragraph()
team.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = team.add_run('Team Members\nMustafa GÖZÜTOK\nMehmet Umur ÖZÜ\nMehmet Yasin TOSUN')
run.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════ TABLE OF CONTENTS ══════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Product Perspective',
    '2. Product Functions',
    '3. User Characteristics',
    '4. Constraints',
    '5. System Features (Use Case Based)',
    '   5.1 Use Case Diagram',
    '   5.2 Use Case 1: Start Game Battle',
    '   5.3 Use Case 2: Execute Combat Actions',
    '   5.4 Use Case 3: Switch Combat Mode',
    '   5.5 Use Case 4: Participate in Local Tournament',
    '6. Non-Functional Requirements',
    '7. External Interface Requirements',
    '8. Software Architecture',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number' if not item.startswith('   ') else 'List Bullet')

doc.add_page_break()

# ══════════════════════ 1. PRODUCT PERSPECTIVE ══════════════════════
doc.add_heading('1. Product Perspective', level=1)

doc.add_paragraph(
    'Color Clash is a standalone web-based two-player fighting game developed as an educational '
    'project for the TOBB ETU BIL 482 course. The application demonstrates the effective use of '
    'object-oriented software design patterns within an interactive game development context. '
    'The game runs entirely in the browser using HTML5 Canvas and TypeScript, requiring no '
    'server-side components or external dependencies beyond the build toolchain (Vite).'
)

doc.add_heading('Context', level=2)
add_table(doc, ['Aspect', 'Description'], [
    ['Platform', 'Web browser application (desktop/laptop focused)'],
    ['Project Type', 'Educational demonstration of design patterns in game development'],
    ['Deployment', 'Self-hosted static web application'],
    ['Technology Stack', 'TypeScript, HTML5 Canvas, Vite build system'],
])

doc.add_heading('Key Characteristics', level=2)
chars = [
    'Local multiplayer game with two players sharing a single keyboard',
    'Real-time combat mechanics with visual feedback',
    'Local tournament mode for group play with multiple participants',
    'Pattern-driven architecture showcasing professional software design practices',
    'No persistence layer – sessions are ephemeral',
    'Child-friendly, non-violent, and colorful game design',
]
for c in chars:
    doc.add_paragraph(c, style='List Bullet')

# ══════════════════════ 2. PRODUCT FUNCTIONS ══════════════════════
doc.add_heading('2. Product Functions', level=1)
doc.add_paragraph('The system provides the following core functions:')

functions = [
    ('Two-Player Combat System',
     'Simultaneous control of two characters on a shared arena with real-time collision '
     'detection, hit registration, health tracking, and victory condition detection.'),
    ('Character-Specific Combat Mode System',
     'A scalable strategy system where each character possesses unique fighting modes '
     '(e.g., Fire/Water for one character, Earth/Wind for another). Players can dynamically switch '
     'between these modes during combat to adapt their strategy, with each mode offering '
     'distinct attributes such as damage output, defense multipliers, movement speed, '
     'and visual effects.'),
    ('Character State Management',
     'Idle, Move, Attack, and Hit states with smooth transitions. Physics-based movement '
     'with gravity and boundaries. Attack cooldowns and hit stun mechanics.'),
    ('Input Handling',
     'Dual keyboard input for two players (WASD/Arrows + action keys). Command pattern '
     'for flexible input mapping. Prevention of input conflicts between players.'),
    ('Visual Feedback System',
     'Particle effects for attacks (themed per mode), health bars with damage animations, '
     'mode indicators showing current stance, and victory screen overlay.'),
    ('Game Flow Control',
     'Pause/resume functionality, restart mechanism, and frame-rate independent game loop.'),
    ('Local Tournament Mode',
     'Registration of multiple players by name, automatic match scheduling with elimination '
     'bracket structure, sequential match play, and progression through rounds to a final match.'),
]
for title, desc in functions:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ══════════════════════ 3. USER CHARACTERISTICS ══════════════════════
doc.add_heading('3. User Characteristics', level=1)

doc.add_heading('Primary Users', level=2)
doc.add_paragraph(
    'Casual gamers, children, young players, and university students interested in local '
    'multiplayer party games, understanding game development patterns, or quick accessible gaming sessions.'
)

doc.add_heading('User Skill Levels', level=2)
add_table(doc, ['User Type', 'Characteristics', 'Needs'], [
    ['Casual Player', 'Familiar with basic keyboard controls, no technical knowledge required',
     'Clear visual feedback, intuitive controls, minimal learning curve'],
    ['Children / Young Players', 'Seeking fun, non-violent fighting game',
     'Child-friendly design, simple mechanics, colorful visuals'],
    ['CS Student / Developer', 'Programming background, interested in design patterns',
     'Well-documented code, clear pattern implementations, extensible architecture'],
    ['Game Enthusiast', 'Experienced with fighting games',
     'Responsive controls, balanced gameplay, strategic depth'],
])

doc.add_heading('Accessibility Considerations', level=2)
items = [
    'No color-blindness accommodations currently (uses distinct color coding per mode)',
    'Two-player local only (requires physical presence)',
    'Keyboard-only input (no gamepad/touch support)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ══════════════════════ 4. CONSTRAINTS ══════════════════════
doc.add_heading('4. Constraints', level=1)

doc.add_heading('Platform Constraints', level=2)
add_table(doc, ['Constraint', 'Description'], [
    ['Browser Compatibility', 'Requires modern browsers supporting ES6+ and HTML5 Canvas (Chrome 90+, Firefox 88+, Edge 90+, Safari 14+)'],
    ['Display Requirements', 'Minimum 1024×576 pixels viewport'],
    ['Input Device', 'Physical keyboard required (mobile/touch not supported)'],
])

doc.add_heading('Technical Constraints', level=2)
add_table(doc, ['Constraint', 'Description'], [
    ['Language', 'TypeScript (strict mode enabled)'],
    ['Build System', 'Vite (specified in package.json)'],
    ['No Backend', 'Pure frontend application, no server communication'],
    ['No Persistence', 'Game state is not saved between sessions'],
    ['Zero Production Dependencies', 'No game engines, no UI frameworks, no utility libraries – pure TypeScript + Browser APIs'],
])

doc.add_heading('Development Constraints', level=2)
add_table(doc, ['Constraint', 'Description'], [
    ['Academic Timeline', 'Designed for BIL 482 course project submission'],
    ['Pattern Requirements', 'Must demonstrate at least 5 design patterns (Strategy, State, Observer, Command, Factory, Singleton, Object Pool)'],
    ['Code Quality', 'Must maintain clean architecture principles'],
    ['Team Size', 'Three-person development team'],
])

doc.add_heading('Performance Constraints', level=2)
add_table(doc, ['Constraint', 'Target'], [
    ['Frame Rate', '60 FPS on standard hardware'],
    ['Input Latency', '< 16ms (1 frame)'],
    ['Memory', 'Efficient particle pooling to prevent memory leaks'],
])

# ══════════════════════ 5. SYSTEM FEATURES ══════════════════════
doc.add_heading('5. System Features (Use Case Based)', level=1)

# 5.1 Use Case Diagram
doc.add_heading('5.1 Use Case Diagram', level=2)
doc.add_paragraph(
    'The following UML Use Case Diagram represents the primary system functionality. '
    'The system has two main actors (Player 1 and Player 2) who interact with four core use cases.'
)

doc.add_paragraph(
    'Actors: Player 1, Player 2\n\n'
    'Use Cases:\n'
    '  • UC1: Start Game Battle – Both players initiate a game session\n'
    '  • UC2: Execute Combat Actions – Players perform movement, attacks, and defense\n'
    '  • UC3: Switch Combat Mode – Players toggle between their character\'s specific combat modes\n'
    '  • UC4: Participate in Local Tournament – Players register and compete in a tournament bracket\n\n'
    'Included Sub-Systems:\n'
    '  • Collision Detection System (included by UC2)\n'
    '  • Particle Effect System (included by UC2)\n'
    '  • Observer Notification System (included by UC3)\n'
    '  • Match Scheduling System (included by UC4)'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[Use Case Diagram – see Appendix or render from Mermaid source below]')
run.italic = True
run.font.size = Pt(9)

doc.add_paragraph(
    'Mermaid Source for rendering:\n'
    'graph TB\n'
    '    %% Actors\n'
    '    P1[Player 1]\n'
    '    P2[Player 2]\n'
    '    TO[Tournament Organizer]\n\n'
    '    %% Use Cases\n'
    '    UC1((Start Game Battle))\n'
    '    UC2((Execute Combat Actions))\n'
    '    UC3((Switch Combat Mode))\n'
    '    UC4((Participate in Local Tournament))\n\n'
    '    %% Sub-Systems\n'
    '    SubCol((Collision Detection))\n'
    '    SubPart((Particle Effects))\n'
    '    SubObs((Observer Notification))\n'
    '    SubSch((Match Scheduling))\n\n'
    '    %% Direct Relations\n'
    '    P1 --> UC1\n'
    '    P2 --> UC1\n'
    '    P1 --> UC2\n'
    '    P2 --> UC2\n'
    '    P1 --> UC3\n'
    '    P2 --> UC3\n'
    '    P1 --> UC4\n'
    '    P2 --> UC4\n'
    '    TO --> UC4\n\n'
    '    %% Dependencies\n'
    '    UC2 -. <<include>> .-> SubCol\n'
    '    UC2 -. <<include>> .-> SubPart\n'
    '    UC3 -. <<include>> .-> SubObs\n'
    '    UC4 -. <<include>> .-> SubSch\n'
    '    UC4 -- <<extends>> --> UC1',
    style='No Spacing'
)

# ── UC1 ──
doc.add_heading('5.2 Use Case 1: Start Game Battle', level=2)
add_table(doc, ['Field', 'Description'], [
    ['Title', 'Start Game Battle'],
    ['Main Actor', 'Player 1 and Player 2'],
    ['Goal', 'Initialize and begin a fighting game session with two characters ready for combat'],
])

doc.add_heading('Preconditions', level=3)
for pc in [
    'User has opened the application in a web browser',
    'Browser supports HTML5 Canvas',
    'Canvas element is successfully rendered',
    'Both players have access to the keyboard',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Main Flow', level=3)
steps_uc1 = [
    'User navigates to the game URL in their browser.',
    'System loads HTML page and executes main.ts entry point.',
    'GameEngine singleton is instantiated with canvas reference.',
    'CharacterFactory creates two Character instances (e.g., FireMage vs WaterGuardian, or default characters with specific modes).',
    'Characters are initialized in their Primary Mode.',
    'InputHandler registers keyboard bindings: Player 1 (WASD + Action Keys), Player 2 (Arrow keys + Action Keys).',
    'GameHUD subscribes to both characters as Observer.',
    'CollisionSystem initializes hitbox tracking.',
    'ParticleSystem creates object pool for visual effects.',
    'Game loop starts at 60 FPS.',
    'System renders initial game state with both characters in idle state.',
    'Control instructions are displayed on screen.',
    'Both players can now input commands.',
]
for i, s in enumerate(steps_uc1, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Postconditions', level=3)
for pc in [
    'Both characters are alive with full health (100 HP)',
    'Characters are in Idle state',
    'Game is in "Playing" state',
    'All systems are actively updating',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Alternative / Exception Flows', level=3)
doc.add_paragraph('Canvas not found: System logs error and game does not initialize.')
doc.add_paragraph('Browser incompatibility: Canvas rendering fails; user is alerted to upgrade browser.')

# ── UC2 ──
doc.add_heading('5.3 Use Case 2: Execute Combat Actions', level=2)
add_table(doc, ['Field', 'Description'], [
    ['Title', 'Execute Combat Actions'],
    ['Main Actor', 'Player 1 or Player 2 (initiator)'],
    ['Goal', 'Perform combat maneuvers including movement, attacking, and dealing/receiving damage'],
])

doc.add_heading('Preconditions', level=3)
for pc in [
    'Game is in "Playing" state',
    'Character is not defeated (health > 0)',
    'Character is not in hit stun state (for attack actions)',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Main Flow – Movement Scenario', level=3)
steps = [
    'Player presses movement key (W/A/S/D or Arrow keys).',
    'InputHandler captures keydown event and creates MoveCommand.',
    'MoveCommand is executed; character inputFlags are updated.',
    'Current state checks for transition; character transitions to MoveState.',
    'MoveState.update() applies velocity based on input direction.',
    'Character.update() applies gravity and checks boundaries.',
    'Character position is updated and rendered at new position.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Main Flow – Attack Scenario', level=3)
steps = [
    'Player presses attack key (F or K).',
    'InputHandler creates AttackCommand; checks cooldown and stun state.',
    'Character transitions to AttackState; attack hitbox is created using current active mode.',
    'CollisionSystem checks for hitbox intersection with opponent.',
    'If hit detected: attacker\'s activeStrategy.attack() calculates damage based on mode multipliers.',
    'Defender\'s activeStrategy.defend() calculates damage reduction (specific to their current mode).',
    'Opponent\'s health is reduced; opponent transitions to HitState.',
    'ParticleSystem spawns character-specific particles at hit location.',
    'GameHUD HealthBar is notified via Observer pattern; health bar animates damage.',
    'Attacker\'s cooldown starts; attacker returns to Idle after attack duration.',
    'Defender exits HitState after hit stun duration expires.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Postconditions', level=3)
for pc in [
    'Character position is updated (for movement)',
    'Opponent health is reduced (if attack connected)',
    'Visual feedback displayed (particles, health bar animation)',
    'Attack cooldown and hit stun timers are active',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Alternative / Exception Flows', level=3)
doc.add_paragraph('Attack on cooldown: No state transition; character remains in current state.')
doc.add_paragraph('Attack misses: No damage dealt; attack animation and cooldown still play.')
doc.add_paragraph('Boundary reached: Position clamped to canvas bounds; velocity zeroed.')
doc.add_paragraph('Fatal damage: Opponent health reaches 0; victory condition triggered.')

# ── UC3 ──
doc.add_heading('5.4 Use Case 3: Switch Combat Mode', level=2)
add_table(doc, ['Field', 'Description'], [
    ['Title', 'Switch Combat Mode'],
    ['Main Actor', 'Player 1 or Player 2'],
    ['Goal', 'Toggle character\'s combat strategy between available modes (e.g., Attack Focus ↔ Defense Focus)'],
])

doc.add_heading('Preconditions', level=3)
for pc in [
    'Game is in "Playing" state',
    'Character is not currently attacking',
    'Character is not in hit stun',
    'Player has not held the switch key (single press detection)',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Main Flow', level=3)
steps = [
    'Player presses switch mode key (G for Player 1, L for Player 2).',
    'InputHandler detects keydown and checks switchModePressed flag to prevent repeat.',
    'SwitchModeCommand is created and executed.',
    'Character checks preconditions (not attacking, not stunned).',
    'Character toggles active strategy reference to the next available mode (e.g., Mode A → Mode B).',
    'New mode properties take effect immediately (Updated damage/defense/speed multipliers).',
    'Character notifies all Observers via Observer pattern with ModeChangeEvent.',
    'GameHUD updates mode indicator (e.g., changing icon from 🔥 to 💧 or �️ to ⚔️).',
    'Character visual rendering updates (glow color, primary color).',
    'On next attack/defense, new mode statistics are applied.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Postconditions', level=3)
for pc in [
    'Character\'s activeStrategy reference points to new mode',
    'Combat statistics changed according to new mode',
    'Visual appearance reflects new mode',
    'GameHUD displays updated mode indicator',
    'All Observers have been notified',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Alternative / Exception Flows', level=3)
doc.add_paragraph('Character is attacking or hit stunned: Mode switch exits early; no change occurs.')
doc.add_paragraph('Key held down: InputHandler prevents repeated toggling.')

# ── UC4 ──
doc.add_heading('5.5 Use Case 4: Participate in Local Tournament', level=2)
add_table(doc, ['Field', 'Description'], [
    ['Title', 'Participate in Local Tournament'],
    ['Main Actor', 'Players (2 or more participants)'],
    ['Goal', 'Register participants and compete through a structured elimination bracket until a tournament winner is determined'],
])

p = doc.add_paragraph()
run = p.add_run('Note: ')
run.bold = True
p.add_run(
    'Tournament mode is a planned feature defined in the Project Definition Document. '
    'It is currently under development and will be implemented as part of the project roadmap.'
)

doc.add_heading('Preconditions', level=3)
for pc in [
    'Application is loaded in the browser',
    'At least 2 players are available to participate',
    'Players have access to the keyboard',
    'Game is in "Menu" state (tournament not yet started)',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Main Flow', level=3)
steps = [
    'Players select "Tournament Mode" from the main menu.',
    'System displays the player registration screen.',
    'Each participant enters their name into the registration form.',
    'Once all players are registered, the organizer confirms the participant list.',
    'System generates a tournament bracket (elimination format) based on the number of participants.',
    'System displays the bracket and schedules the first match.',
    'The two players for the current match take positions at the keyboard.',
    'A standard game battle is initiated (see UC1: Start Game Battle).',
    'Players compete using combat actions (see UC2) and mode switching (see UC3).',
    'When a player wins the match, the system records the result and updates the bracket.',
    'System advances the winner to the next round.',
    'Steps 7–11 repeat for each scheduled match until the final match is played.',
    'System declares the tournament winner and displays the final results screen.',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('Postconditions', level=3)
for pc in [
    'A tournament winner has been declared',
    'The complete bracket with all match results is displayed',
    'Players can choose to start a new tournament or return to the main menu',
]:
    doc.add_paragraph(pc, style='List Bullet')

doc.add_heading('Alternative / Exception Flows', level=3)
doc.add_paragraph('Fewer than 2 players: System requires a minimum of 2 participants to start.')
doc.add_paragraph('Player disconnects (closes tab): Match is forfeited; opponent advances.')
doc.add_paragraph('Odd number of players: One player receives a bye and advances automatically to the next round.')

# ══════════════════════ 6. NON-FUNCTIONAL REQUIREMENTS ══════════════════════
doc.add_heading('6. Non-Functional Requirements', level=1)

doc.add_heading('6.1 Usability', level=2)
add_table(doc, ['Requirement', 'Description', 'Metric'], [
    ['Learning Curve', 'New players should understand basic controls within 1 minute', 'Time to first intentional attack < 60s'],
    ['Visual Clarity', 'Health bars, mode indicators, and effects must be immediately recognizable', '90% can identify mode/health without instruction'],
    ['Control Responsiveness', 'Input commands should feel instantaneous', 'Input-to-visual latency < 16ms (1 frame)'],
    ['Instructions', 'On-screen control guide visible at all times', 'Keyboard shortcuts displayed in HUD'],
    ['Error Prevention', 'No invalid state transitions possible', 'State machine prevents illegal transitions'],
])

doc.add_heading('6.2 Performance', level=2)
add_table(doc, ['Requirement', 'Target', 'Acceptance Criteria'], [
    ['Frame Rate', '60 FPS', 'Maintain ≥55 FPS on mid-range hardware'],
    ['Input Latency', '< 16ms', 'Keypress to state change < 1 frame'],
    ['Particle Count', '100+ particles', 'No FPS drop below 50 with max particle load'],
    ['Memory Usage', '< 100MB heap', 'No memory leaks during 30+ min play'],
    ['Load Time', '< 2 seconds', 'Page load to playable state'],
    ['Delta Time', 'Frame-rate independent', 'Identical behavior at 30–120 FPS'],
])

doc.add_heading('6.3 Portability', level=2)
add_table(doc, ['Platform', 'Support Level', 'Notes'], [
    ['Windows', 'Fully Supported', 'Primary development platform'],
    ['macOS', 'Fully Supported', 'Tested on Safari and Chrome'],
    ['Linux', 'Fully Supported', 'Tested on Firefox and Chromium'],
    ['Mobile/Tablet', 'Not Supported', 'No touch controls, keyboard required'],
])
doc.add_paragraph('Browser Compatibility: Chrome 90+, Firefox 88+, Edge 90+, Safari 14+. Internet Explorer is not supported.')

doc.add_heading('6.4 Reliability', level=2)
add_table(doc, ['Aspect', 'Requirement', 'Measure'], [
    ['Crash Resistance', 'No crashes during normal gameplay', 'MTBF > 10 hours'],
    ['State Consistency', 'Game state always valid', 'All state transitions validated'],
    ['Error Handling', 'Graceful degradation', 'Try-catch in initialization, console logging'],
    ['Recovery', 'Restart always works', 'R key restart succeeds 100% from GameOver'],
    ['Tab Switching', 'Survives tab switch', 'Game pauses on tab blur event'],
    ['Memory Leaks', 'No accumulating memory', 'Heap stable after 1000 particles spawned/destroyed'],
])

# ══════════════════════ 7. EXTERNAL INTERFACE REQUIREMENTS ══════════════════════
doc.add_heading('7. External Interface Requirements', level=1)

doc.add_heading('7.1 User Interfaces', level=2)

doc.add_heading('Main Game Canvas', level=3)
items = [
    'Dimensions: 1024×576 pixels (16:9 aspect ratio)',
    'Rendering: HTML5 Canvas 2D context',
    'Background: Gradient from dark purple (#2d1b4e) to dark blue (#1a1a2e)',
    'Arena: Ground line at y=480 with visual decorations',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('HUD Elements', level=3)
add_table(doc, ['Element', 'Position', 'Description'], [
    ['Player 1 Health Bar', 'Top-left', 'Red bar, 300px width, shows damage animations'],
    ['Player 2 Health Bar', 'Top-right', 'Blue bar, 300px width, mirrored from P1'],
    ['Player 1 Mode Indicator', 'Below P1 health bar', 'Circle with Fire/Water emoji and color'],
    ['Player 2 Mode Indicator', 'Below P2 health bar', 'Circle with Fire/Water emoji and color'],
    ['Center Title', 'Top-center', '"⚔️ FIGHT ⚔️"'],
    ['Control Instructions', 'Bottom-center', 'Keyboard shortcuts, small font'],
])

doc.add_heading('Victory Screen Overlay', level=3)
doc.add_paragraph(
    'Triggered when a player\'s health reaches 0. Displays a semi-transparent black overlay '
    'with winner announcement text, player color theme, and glowing effects.'
)

doc.add_heading('Tournament Screens (Planned)', level=3)
items = [
    'Player Registration Screen: Input fields for participant names with add/remove functionality.',
    'Tournament Bracket View: Visual bracket showing all scheduled matches and results.',
    'Match Result Screen: Displays winner of each match with option to proceed to next match.',
    'Tournament Winner Screen: Final celebration screen displaying the overall tournament champion.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.2 Software Interfaces', level=2)

doc.add_heading('Browser APIs', level=3)
add_table(doc, ['API', 'Purpose', 'Usage'], [
    ['Canvas 2D Context', 'Rendering', 'canvas.getContext("2d") for all drawing operations'],
    ['requestAnimationFrame', 'Game loop', 'Smooth 60 FPS updates synced with display refresh'],
    ['KeyboardEvent', 'Input', 'keydown/keyup listeners for player controls'],
    ['Console API', 'Debugging', 'Logging startup info, errors, warnings'],
])

doc.add_heading('Design Pattern Interfaces', level=3)
add_table(doc, ['Pattern', 'Interface', 'Location'], [
    ['Strategy', 'IModeStrategy', '/src/patterns/strategy/IModeStrategy.ts'],
    ['State', 'ICharacterState', '/src/patterns/state/ICharacterState.ts'],
    ['Command', 'ICommand', '/src/patterns/command/Command.ts'],
    ['Observer', 'IObserver<T>, ISubject<T>', '/src/patterns/observer/Observer.ts'],
    ['Factory', 'CharacterFactory', '/src/patterns/factory/CharacterFactory.ts'],
    ['Object Pool', 'ObjectPool<T>', '/src/patterns/pool/ObjectPool.ts'],
])

doc.add_heading('Build System Interface', level=3)
doc.add_paragraph('Vite Configuration: Development server on port 5173 with HMR enabled, TypeScript compilation integrated.')
doc.add_paragraph('Package Scripts: "dev" (vite), "build" (tsc && vite build), "preview" (vite preview).')
doc.add_paragraph('DevDependencies: typescript ~5.6.2, vite ^6.0.5, @types/node ^22.0.0. Zero production dependencies.')

# ══════════════════════ 8. SOFTWARE ARCHITECTURE ══════════════════════
doc.add_heading('8. Software Architecture', level=1)

doc.add_heading('8.1 High-Level Architecture', level=2)
doc.add_paragraph(
    'Color Clash follows a layered monolithic architecture with clear separation of concerns '
    'organized by design patterns. The architecture is structured as a single-page application '
    'with a pattern-driven game engine core.'
)

doc.add_heading('Architectural Layers', level=3)
add_table(doc, ['Layer', 'Responsibility', 'Key Components'], [
    ['Presentation Layer', 'User interface and visual rendering', 'HTML5 Canvas, GameHUD, HealthBar'],
    ['Game Engine Core', 'Central orchestration (Singleton)', 'GameEngine, Game Loop, Renderer'],
    ['Game Systems Layer', 'Domain-specific subsystems', 'InputHandler (Command), CollisionSystem, ParticleSystem (Object Pool)'],
    ['Entity Layer', 'Core game objects', 'Character (Strategy+State+Observer), Hitbox, Particle'],
    ['Pattern Implementations', 'Reusable design pattern abstractions', 'Strategy, State, Observer, Command, Factory, Object Pool'],
])

doc.add_heading('Architecture Diagram Description', level=3)
doc.add_paragraph(
    'The system uses a top-down layered architecture:\n\n'
    '1. Presentation Layer: HTML5 Canvas UI → GameHUD → HealthBar Components\n'
    '2. Game Engine Core (Singleton): GameEngine → Game Loop (60 FPS) → Rendering System\n'
    '3. Game Systems Layer: InputHandler (Command Pattern) | CollisionSystem | ParticleSystem (Object Pool)\n'
    '4. Entity Layer: Character Entities (Strategy + State) | Hitbox | Particle\n'
    '5. Pattern Implementations: Strategy (Mode A/Mode B) | State (Idle/Move/Attack/Hit) | Observer | Command | Factory | Pool\n\n'
    'Communication flows: HTML loads Engine → Engine drives Game Loop → Loop calls Input → Characters → Particles → Renderer. '
    'Input creates Commands. Characters use Strategy and State patterns. Characters notify Observers (HUD/HealthBar). '
    'Engine creates Characters via Factory. ParticleSystem uses Object Pool.'
)

doc.add_heading('8.2 How Architecture Facilitates Use Cases', level=2)

uc_arch = [
    ('UC1: Start Game Battle',
     'Singleton GameEngine ensures single instance initialization. Factory Pattern creates '
     'diverse character instances. Observer Pattern connects UI to entities. InputHandler '
     'registers bindings. Object Pool pre-allocates particles.'),
    ('UC2: Execute Combat Actions',
     'Command Pattern translates keypresses into actions. State Pattern manages character transitions. '
     'Strategy Pattern applies mode-specific damage/defense (extensible for multiple characters). '
     'CollisionSystem detects hits. Observer notifies UI.'),
    ('UC3: Switch Combat Mode',
     'Strategy Pattern allows dynamic switching of character-specific combat behaviors. '
     'Command Pattern handles switch input. Observer Pattern notifies HUD. '
     'State Machine validates mode switch is allowed.'),
    ('UC4: Participate in Local Tournament',
     'Tournament Manager orchestrates the bracket and match scheduling. Factory Pattern creates '
     'fresh Character instances per match. GameEngine restart mechanism resets state between matches. '
     'Observer Pattern notifies tournament UI of match results.'),
]
for title, desc in uc_arch:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('8.3 Architectural Benefits', level=2)
add_table(doc, ['Design Goal', 'Architectural Solution'], [
    ['Modularity', 'Layered architecture with clear interfaces between layers'],
    ['Extensibility', 'Strategy pattern allows unlimited new character modes without changing Character class'],
    ['Testability', 'Decoupled systems, dependency injection of observers'],
    ['Performance', 'Object pooling, efficient collision detection, delta-time game loop'],
    ['Maintainability', 'Pattern-based organization, single responsibility classes'],
    ['Scalability', 'Easy to add new characters, states, commands, or particle types'],
])

doc.add_heading('8.4 Design Patterns Summary', level=2)
add_table(doc, ['Pattern', 'Usage', 'Location', 'Purpose'], [
    ['Singleton', '1 instance', 'GameEngine', 'Ensure single game orchestrator'],
    ['Strategy', 'Multiple strategies', 'IModeStrategy implementations', 'Define unique combat behaviors per mode/character'],
    ['State', '4 states', 'IdleState, MoveState, AttackState, HitState', 'Manage character behavior transitions'],
    ['Observer', '2 observers', 'HealthBar, GameHUD', 'Decouple UI from game logic'],
    ['Command', '3 commands', 'MoveCommand, AttackCommand, SwitchModeCommand', 'Encapsulate input actions'],
    ['Factory', '1 factory', 'CharacterFactory', 'Centralize character creation with specific configs'],
    ['Object Pool', '1 pool', 'ParticleSystem for Particle', 'Reuse particles for performance'],
])


# ══════════════════════ 9. TASK ASSIGNMENT MATRIX ══════════════════════
doc.add_page_break()
doc.add_heading('9. Task Assignment Matrix', level=1)

doc.add_paragraph(
    'The following matrix outlines the distribution of responsibilities among team members '
    'for the key components and phases of the project.'
)

# Headers: SRD Section / Task, Members
headers = ['SRD Section / Contribution', 'Mustafa GÖZÜTOK', 'Mehmet Umur ÖZÜ', 'Mehmet Yasin TOSUN']
rows = [
    # General Sections
    ['1. Product Perspective & Introduction', 'X', '', ''],
    ['2. Product Functions Overview', 'X', '', ''],
    ['3. User Characteristics', '', '', 'X'],
    ['4. Constraints (Technical & Dev)', '', 'X', ''],
    
    # Use Cases
    ['5. Use Case Definitions (UC1-UC4)', 'X', '', 'X'],
    ['5.1 Use Case Diagram (Design)', '', 'X', ''],
    
    # Technical Requirements
    ['6. Non-Functional Requirements', '', 'X', ''],
    ['7. External Interface Requirements', '', '', 'X'],
    
    # Architecture
    ['8. Software Architecture & Patterns', '', 'X', ''],
    
    # Document Management
    ['Review & Final Edits', 'X', 'X', 'X'],
]

table = doc.add_table(rows=1 + len(rows), cols=len(headers))
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set headers
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True

# Fill data
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        cell = table.rows[ri + 1].cells[ci]
        cell.text = val
        if val == 'X':
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            if len(cell.paragraphs[0].runs) > 0:
                 run.bold = True

doc.add_paragraph()
doc.add_paragraph('X: Principal Contributor')


# ══════════════════════ SAVE ══════════════════════
output_path = 'ColorClash-Software-Requirements-Document-v5.docx'
doc.save(output_path)
print(f'SUCCESS: Document saved to {output_path}')

